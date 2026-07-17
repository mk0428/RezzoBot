"""Resume exporter — PDF and DOCX generation from resume text"""
import io
import os
import logging
import tempfile
from typing import Optional

logger = logging.getLogger(__name__)

# HTML template for PDF export
PDF_TEMPLATE = """<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  @page {{
    size: letter;
    margin: 0.7in 0.8in;
  }}
  body {{
    font-family: 'Helvetica Neue', 'Segoe UI', Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.5;
    color: #222;
  }}
  h1 {{
    font-size: 22pt;
    font-weight: 700;
    margin-bottom: 2px;
    color: #000;
  }}
  .contact {{
    font-size: 10pt;
    color: #555;
    margin-bottom: 16px;
  }}
  h2 {{
    font-size: 12pt;
    font-weight: 700;
    text-transform: uppercase;
    letter-spacing: 1px;
    color: #333;
    border-bottom: 2px solid #333;
    padding-bottom: 3px;
    margin-top: 16px;
    margin-bottom: 8px;
  }}
  .section {{
    margin-bottom: 12px;
  }}
  .exp-header {{
    display: flex;
    justify-content: space-between;
    font-weight: 600;
    font-size: 11pt;
    margin-bottom: 2px;
  }}
  .exp-company {{
    color: #000;
  }}
  .exp-date {{
    color: #555;
    font-weight: 400;
  }}
  .exp-title {{
    font-size: 10.5pt;
    color: #444;
    font-style: italic;
    margin-bottom: 4px;
  }}
  ul {{
    margin: 4px 0 8px 0;
    padding-left: 18px;
  }}
  li {{
    font-size: 10.5pt;
    margin-bottom: 2px;
    color: #333;
  }}
  .skills {{
    font-size: 10.5pt;
    color: #333;
  }}
  .education-item {{
    margin-bottom: 4px;
  }}
  .edu-header {{
    display: flex;
    justify-content: space-between;
    font-weight: 600;
    font-size: 10.5pt;
  }}
  .edu-detail {{
    font-size: 10pt;
    color: #555;
  }}
</style>
</head>
<body>
{content}
</body>
</html>"""


def _resume_text_to_html(resume_text: str) -> str:
    """Convert plain text resume to styled HTML."""
    lines = resume_text.split("\n")
    html_parts = []
    in_list = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            continue

        # Remove markdown bold markers for display
        clean = stripped.replace("**", "")

        # Name (first non-empty line, or any line with no bold/section markers)
        if len(html_parts) == 0:
            html_parts.append(f"<h1>{clean}</h1>")
            continue

        # Contact info (lines with email/linkedin/phone symbols)
        if any(c in clean for c in ["@", "linkedin", "linkedin.com"]):
            html_parts.append(f'<div class="contact">{clean}</div>')
            continue

        # Section headers (ALL CAPS lines, possibly with markdown bold)
        if stripped.replace("*", "").strip().isupper() and len(clean) < 60 and not clean.startswith("•") and not clean.startswith("-"):
            if in_list:
                html_parts.append("</ul>")
                in_list = False
            html_parts.append(f"<h2>{clean}</h2>")
            continue

        # Bold headers (lines with company/date pattern)
        if "|" in clean and len(clean) < 120:
            parts = [p.strip() for p in clean.split("|")]
            html_parts.append(f"<div class='exp-header'><span class='exp-company'>{' | '.join(parts)}</span></div>")
            continue

        # Bullet points
        if stripped.startswith("•") or stripped.startswith("-") or stripped.startswith("*"):
            text = stripped.lstrip("•-* ").strip().replace("**", "")
            if not in_list:
                html_parts.append("<ul>")
                in_list = True
            html_parts.append(f"<li>{text}</li>")
            continue

        # Regular text
        if in_list:
            html_parts.append("</ul>")
            in_list = False
        html_parts.append(f"<p>{clean}</p>")

    if in_list:
        html_parts.append("</ul>")

    return "\n".join(html_parts)


def export_pdf(resume_text: str) -> Optional[bytes]:
    """Export resume text as PDF bytes."""
    try:
        from weasyprint import HTML

        body_html = _resume_text_to_html(resume_text)
        full_html = PDF_TEMPLATE.format(content=body_html)

        pdf_bytes = HTML(string=full_html).write_pdf()
        return pdf_bytes

    except ImportError:
        logger.error("WeasyPrint not installed")
        return None
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        return None


def export_docx(resume_text: str) -> Optional[bytes]:
    """Export resume text as DOCX bytes."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        lines = resume_text.split("\n")
        in_list = False

        for line in lines:
            stripped = line.strip()
            if not stripped:
                if in_list:
                    in_list = False
                continue

            # Name
            if len(doc.paragraphs) <= 1:
                p = doc.add_paragraph()
                p.alignment = 0  # LEFT alignment
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(22)
                continue

            # Contact
            if any(c in stripped for c in ["@", "linkedin", "linkedin.com"]):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                continue

            # Section headers
            if stripped.isupper() and len(stripped) < 50 and not stripped.startswith("•"):
                if in_list:
                    in_list = False
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)

                # Underline
                p_border = p.paragraph_format._element
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr")
                bottom = OxmlElement("w:bottom")
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "4")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "333333")
                pBdr.append(bottom)
                pPr.append(pBdr)
                continue

            # Company headers with dates
            if "|" in stripped and len(stripped) < 120:
                if in_list:
                    in_list = False
                parts = [p.strip() for p in stripped.split("|")]
                p = doc.add_paragraph()
                run = p.add_run(" | ".join(parts))
                run.bold = True
                run.font.size = Pt(11)
                continue

            # Bullet points
            if stripped.startswith("•") or stripped.startswith("-") or stripped.startswith("*"):
                text = stripped.lstrip("•-* ").strip()
                p = doc.add_paragraph(text, style="List Bullet")
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                in_list = True
                continue

            # Regular text
            if in_list:
                in_list = False
            p = doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    except ImportError:
        logger.error("python-docx not installed")
        return None
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return None
